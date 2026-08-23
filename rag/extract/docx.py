"""DOCX, via python-docx.

Two things need care.

**Order.** `document.paragraphs` and `document.tables` are separate lists, and
reading one then the other puts every table at the end, detached from the text
that introduced it. The body's XML children are walked instead, so a table
stays where the author put it.

**Headings.** Word marks them with a paragraph style, not with markup, and the
style name is localised and freely renamed. So the style is checked for the
`Heading N` pattern and also for an outline level, and anything unrecognised
stays ordinary text rather than being guessed at.

DOCX has no pages. Word computes them at layout time and the file does not
store them, so every element here carries `page=None` - which is the honest
answer, and keeps `summarize_document` from citing a page that does not exist.
"""

from __future__ import annotations

import re
from pathlib import Path

from rag.elements import DocumentElement, ExtractedDocument, render_table
from rag.extract.clean import clean_text

# "Heading 1", and the localised variants that still carry the digit.
_HEADING_STYLE = re.compile(r"heading\s*(\d+)", re.IGNORECASE)

# Word's own name for the document title.
_TITLE_STYLES = frozenset({"title", "subtitle"})


class DocxError(Exception):
    """A DOCX could not be opened or read."""


def _docx():
    try:
        import docx
    except ImportError:
        raise DocxError(
            "Reading DOCX files needs python-docx. Install it with: "
            "pip install -r requirements-documents.txt"
        ) from None
    return docx


def extract_docx(path: Path) -> ExtractedDocument:
    """Read a DOCX into elements, in document order."""
    docx = _docx()

    try:
        document = docx.Document(str(path))
    except Exception as exc:
        raise DocxError(
            f"{path.name} could not be opened as a DOCX "
            f"({type(exc).__name__}: {exc}). It may be corrupt, or it may be "
            f"an older .doc file, which is a different format."
        ) from None

    try:
        elements = list(_walk(document))
    except DocxError:
        raise
    except Exception as exc:
        raise DocxError(
            f"{path.name} has a structure this could not read "
            f"({type(exc).__name__}: {exc})."
        ) from None

    return ExtractedDocument(elements=elements, kind="docx")


def _walk(document):
    """Yield elements in the order they appear in the body."""
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    body = document.element.body
    # Maps the raw XML children back to the objects python-docx wraps them in.
    for child in body.iterchildren():
        tag = child.tag.rsplit("}", 1)[-1]

        if tag == "p":
            element = _paragraph(Paragraph(child, document))
            if element is not None:
                yield element
        elif tag == "tbl":
            rendered = _table(Table(child, document))
            if rendered:
                yield DocumentElement(type="table", content=rendered)


def _paragraph(paragraph) -> DocumentElement | None:
    text = clean_text(paragraph.text or "")
    if not text.strip():
        return None

    level = _heading_level(paragraph)
    if level is not None:
        return DocumentElement(type="heading", content=text, level=level)
    return DocumentElement(type="text", content=text)


def _heading_level(paragraph) -> int | None:
    """The heading level of a paragraph, or None if it is body text."""
    name = ""
    try:
        name = (paragraph.style.name or "").strip()
    except Exception:
        name = ""

    if name.lower() in _TITLE_STYLES:
        return 1

    match = _HEADING_STYLE.search(name)
    if match:
        try:
            return max(1, min(int(match.group(1)), 9))
        except ValueError:
            return None

    # A renamed style can still carry Word's outline level, which is what the
    # navigation pane uses.
    try:
        outline = paragraph.paragraph_format.element.pPr.outlineLvl
        if outline is not None and outline.val is not None:
            return max(1, min(int(outline.val) + 1, 9))
    except Exception:
        pass
    return None


def _table(table) -> str:
    """Render a table, skipping any that cannot be read."""
    try:
        rows = [[cell.text for cell in row.cells] for row in table.rows]
    except Exception:
        return ""
    return render_table(rows)
